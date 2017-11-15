#!/usr/bin/env python3
# -*- coding:utf-8 -*-

# 批量修改文件名与文件内容



import os
from docx import Document

path = "/Users/taozz/Dropbox/projectCode/python/python_learning/other/logfile"

temp = 0
for file in os.listdir(path):

    # 查找字符串
    if file.find("陈晓彬") != -1:
        # 计数
        temp += 1
        # print(file)

        get_name = (os.path.splitext(file))
        file_name_head = get_name[0]
        file_name_head = file_name_head.replace("陈晓彬", "梁祥桃")  # 替换

        file_name_rear = get_name[-1]

        # new_file_name = os.path.join(".", (file_name_head + file_name_rear))
        new_file_name = file_name_head + file_name_rear
        # print(new_file_name)
        try:
            # 文件重名名
            # os.rename(file, new_file_name)
            document = Document(file)
            # 读取文档中所有的段落列表
            tables = document.tables
            tables[0].cell(0, 1).text = "201441410237"
            tables[0].cell(0, 3).text = "14软件工程1班"
            tables[0].cell(1, 1).text = "梁祥桃"
            tables[1].cell(1, 1).text = "项目经理：梁祥桃"
            # 保存是否能覆盖？目测不能
            document.save(new_file_name)

        except IOError:
            print("Error: 处理出错")
        else:
            print("第{}个文件处理成功".format(temp))

        # document = Document(path)
        # 读取文档中所有的段落列表
        # tables = document.tables
        # tables[0].cell(0, 1).text = "201441410237"
        # tables[0].cell(0, 3).text = "14软件工程1班"
        # tables[0].cell(1, 1).text = "梁祥桃"

        # tables[1].cell(1, 1).text = "项目经理：梁祥桃"
        # document.save("1.docx")

    else:
        continue
